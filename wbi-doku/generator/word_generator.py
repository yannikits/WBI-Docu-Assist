"""
Word-Dokument-Generator – WBI WIVIO
"""

import io
import re

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree


# ── Stil-Kandidaten ────────────────────────────────────────────────────────────
_STYLE_CANDIDATES = {
    'heading1': ['Heading 1', 'Überschrift 1'],
    'heading2': ['Heading 2', 'Überschrift 2'],
    'heading3': ['Heading 3', 'Überschrift 3'],
    'normal':   ['Normal', 'Standard', 'Body Text', 'Fließtext'],
    'table':    ['Table Grid', 'Tabellenraster', 'Table Normal', 'Normale Tabelle'],
    'bullet':   ['List Bullet', 'Listenabsatz', 'Aufzählung'],
    'number':   ['List Number', 'Listennummer'],
}

_SKIP_SECTION_KEYWORDS = [
    'inhaltsverzeichnis',
    'table of contents',
    'hinweise zur vorlage',
]


def _build_style_map(doc: Document) -> dict:
    """Ermittelt verfügbare Stilnamen aus der Vorlage."""
    available = {s.name for s in doc.styles}
    mapping = {}
    for key, candidates in _STYLE_CANDIDATES.items():
        for candidate in candidates:
            if candidate in available:
                mapping[key] = candidate
                break
    # H3/H4 auf H2 mappen wenn kein eigener H3-Stil vorhanden
    if 'heading3' not in mapping and 'heading2' in mapping:
        mapping['heading3'] = mapping['heading2']
    if 'heading4' not in mapping and 'heading2' in mapping:
        mapping['heading4'] = mapping['heading2']
    return mapping


def _get_template_heading_color(doc: Document, style_map: dict) -> RGBColor:
    """Liest die Farbe von Heading 1/2 aus der Vorlage aus."""
    for key in ('heading1', 'heading2'):
        style_name = style_map.get(key)
        if style_name:
            try:
                color = doc.styles[style_name].font.color.rgb
                if color:
                    return color
            except Exception:
                pass
    return RGBColor(0x1F, 0x38, 0x64)  # WBI-Dunkelblau als letzter Fallback


def _add_heading_safe(doc: Document, text: str, level: int,
                      style_map: dict, fallback_color: RGBColor):
    """Fügt Überschrift ein – Vorlagen-Stil bevorzugt, direktes Format als Fallback."""
    # H3/H4 → H2 falls kein eigener Stil
    effective_level = level
    if level >= 3 and f'heading{level}' not in _STYLE_CANDIDATES:
        effective_level = 2
    style_name = style_map.get(f'heading{effective_level}') or style_map.get('heading2')

    if style_name:
        try:
            return doc.add_paragraph(text, style=style_name)
        except KeyError:
            pass
    # Fallback
    sizes = {1: 16, 2: 14, 3: 13, 4: 12}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 12))
    run.font.color.rgb = fallback_color
    p.paragraph_format.space_before = Pt(0 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_table_safe(doc: Document, style_map: dict, rows: int, cols: int):
    table = doc.add_table(rows=rows, cols=cols)
    style_name = style_map.get('table')
    if style_name:
        try:
            table.style = style_name
        except KeyError:
            pass
    return table


def _safe_filename(title: str) -> str:
    return re.sub(r'[^\w\s\-äöüÄÖÜß]', '_', title).strip() + ".docx"


def _get_sect_pr(doc: Document):
    """Gibt das sectPr-Element zurück (muss immer letztes Element in body sein)."""
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'sectPr':
            return child
    return None


def _insert_before_sect(doc: Document, element):
    """Fügt ein XML-Element vor dem sectPr ein (korrekte Word-Position)."""
    sect_pr = _get_sect_pr(doc)
    if sect_pr is not None:
        sect_pr.addprevious(element)
    else:
        doc.element.body.append(element)


def _extract_and_clear_body(doc: Document):
    """
    Sichert das TOC-Element (sdt), entfernt alle Inhaltselemente außer sectPr.
    Gibt das gesicherte TOC-Element zurück (oder None).
    """
    body = doc.element.body
    children = list(body)
    toc_element = None

    for child in children:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'sdt':
            xml_str = etree.tostring(child, encoding='unicode')
            if 'TOC' in xml_str or 'Inhaltsverzeichnis' in xml_str:
                toc_element = child
                break

    # Alles außer sectPr entfernen
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag not in ('sectPr',):
            body.remove(child)

    return toc_element


def _add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_warning_box(doc: Document, text: str):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '360')
    pPr.append(ind)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _add_refs_table(doc: Document, refs: list, style_map: dict):
    table = _add_table_safe(doc, style_map, rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "DMS-Link / WBI-Nummer"
    hdr[1].text = "Bezeichnung"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for ref in refs:
        if ref.get('num') or ref.get('name'):
            row = table.add_row().cells
            row[0].text = ref.get('num', '')
            row[1].text = ref.get('name', '')


def _strip_md_links(text: str) -> str:
    """[Text](url) → Text, aber keine Bilder (![...] bleibt unverändert)"""
    return re.sub(r'(?<!!)\[([^\]]+)\]\([^)]*\)', r'\1', text)


def _try_insert_image(doc: Document, line: str) -> bool:
    """
    Prüft ob die Zeile ein Base64-Bild ist (![alt](data:image/...;base64,...)).
    Wenn ja: Bild in Word einfügen, True zurückgeben.
    """
    m = re.match(r'!\[([^\]]*)\]\(data:image/([^;]+);base64,([^)]+)\)', line.strip())
    if not m:
        return False
    alt = m.group(1) or 'Bild'
    b64_data = m.group(3)
    try:
        import base64
        from docx.shared import Cm
        img_bytes = base64.b64decode(b64_data)
        img_buf = io.BytesIO(img_bytes)
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(img_buf, width=Cm(14))
        if alt and alt != 'Bild':
            cap = doc.add_paragraph()
            cap_run = cap.add_run(alt)
            cap_run.italic = True
            cap_run.font.size = Pt(9)
            cap_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    except Exception:
        doc.add_paragraph(f'[Bild: {alt}]')
    return True


def _parse_inline(run_text: str):
    segments = []
    pattern = re.compile(r'(\*\*(.+?)\*\*|_(.+?)_|`(.+?)`)')
    last = 0
    for m in pattern.finditer(run_text):
        if m.start() > last:
            segments.append((run_text[last:m.start()], False, False))
        raw = m.group(0)
        inner = m.group(2) or m.group(3) or m.group(4)
        segments.append((inner, raw.startswith('**'), raw.startswith('_') or raw.startswith('`')))
        last = m.end()
    if last < len(run_text):
        segments.append((run_text[last:], False, False))
    return segments or [(run_text, False, False)]


def _add_paragraph_with_inline(doc: Document, text: str, style_map: dict,
                                style_key: str = 'normal'):
    style_name = style_map.get(style_key) or style_map.get('normal')
    try:
        p = doc.add_paragraph(style=style_name) if style_name else doc.add_paragraph()
    except KeyError:
        p = doc.add_paragraph()
    for seg_text, bold, italic in _parse_inline(text):
        run = p.add_run(seg_text)
        run.bold = bold
        run.italic = italic
    return p


def _is_skip_heading(text: str) -> bool:
    clean = re.sub(r'^\d+\.\s*', '', text).strip().lower()
    return any(kw in clean for kw in _SKIP_SECTION_KEYWORDS)


def _add_markdown_content(doc: Document, markdown: str, style_map: dict,
                           fallback_color: RGBColor, skip_first_h1: bool = False):
    lines = markdown.split('\n')
    i = 0
    first_h1_done = not skip_first_h1
    in_skip_section = False
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            return
        data = [r for r in table_rows if not re.match(r'^\|[-| :]+\|$', r.strip())]
        if not data:
            in_table = False; table_rows = []; return
        parsed = [[c.strip() for c in r.strip().strip('|').split('|')] for r in data]
        if not parsed:
            in_table = False; table_rows = []; return
        max_cols = max(len(r) for r in parsed)
        table = _add_table_safe(doc, style_map, rows=len(parsed), cols=max_cols)
        for ri, row in enumerate(parsed):
            for ci, cell_text in enumerate(row):
                if ci < max_cols:
                    cell = table.cell(ri, ci)
                    cell.text = _strip_md_links(cell_text)
                    if ri == 0:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
        in_table = False; table_rows = []

    while i < len(lines):
        raw_line = lines[i]
        stripped = raw_line.strip()

        # Erste H1 überspringen (Titel bereits explizit eingefügt)
        if not first_h1_done and stripped.startswith('# '):
            first_h1_done = True
            i += 1; continue

        # Skip-Abschnitte (TOC, Hinweise)
        if re.match(r'^#{1,3}\s', stripped):
            heading_text = re.sub(r'^#{1,6}\s+', '', stripped)
            if _is_skip_heading(heading_text):
                in_skip_section = True
                i += 1; continue
            else:
                in_skip_section = False

        if in_skip_section:
            i += 1; continue

        # Bilder (Base64) – vor Link-Bereinigung prüfen
        if raw_line.strip().startswith('!['):
            if _try_insert_image(doc, raw_line):
                i += 1; continue

        # Markdown-Links bereinigen
        line = _strip_md_links(raw_line)
        stripped = line.strip()

        # Tabellen
        if stripped.startswith('|'):
            if not in_table:
                in_table = True; table_rows = []
            table_rows.append(line)
            i += 1; continue
        else:
            if in_table:
                flush_table()

        if not stripped or stripped == '---':
            if stripped == '---':
                _add_horizontal_rule(doc)
            else:
                doc.add_paragraph()
            i += 1; continue

        if stripped.startswith('#### '):
            _add_heading_safe(doc, stripped[5:], 4, style_map, fallback_color)
        elif stripped.startswith('### '):
            _add_heading_safe(doc, re.sub(r'^\d+\.\d+\s+', '', stripped[4:]), 3, style_map, fallback_color)
        elif stripped.startswith('## '):
            _add_heading_safe(doc, re.sub(r'^\d+\.\s+', '', stripped[3:]), 2, style_map, fallback_color)
        elif stripped.startswith('# '):
            _add_heading_safe(doc, stripped[2:], 1, style_map, fallback_color)
        elif stripped.startswith('> '):
            _add_warning_box(doc, stripped[2:])
        elif stripped.startswith('- ') or stripped.startswith('* '):
            _add_paragraph_with_inline(doc, stripped[2:], style_map, 'bullet')
        elif re.match(r'^\d+\.\s', stripped):
            _add_paragraph_with_inline(doc, re.sub(r'^\d+\.\s', '', stripped), style_map, 'number')
        else:
            _add_paragraph_with_inline(doc, stripped, style_map)

        i += 1

    if in_table:
        flush_table()


def generate_word(data: dict, template_path: str):
    """Erstellt ein Word-Dokument basierend auf der WBI-Vorlage."""
    title = data.get('titleSubject', '') + ' - ' + data.get('titleTopic', '')
    title = title.strip(' -')
    aushang = data.get('aushang', False)
    chapters = data.get('chapters', [])
    refs = [r for r in data.get('refs', []) if r.get('num') or r.get('name')]
    markdown_content = data.get('markdownContent', '').strip()
    template_id = data.get('template', 'intern')

    doc = Document(template_path)
    style_map = _build_style_map(doc)
    fallback_color = _get_template_heading_color(doc, style_map)

    # Body leeren, TOC sichern
    toc_element = _extract_and_clear_body(doc)

    # ── 1. Titel ───────────────────────────────────────────────────────────────
    _add_heading_safe(doc, title, 1, style_map, fallback_color)

    # ── 2. TOC aus Vorlage wiederherstellen (vor sectPr einfügen!) ─────────────
    if toc_element is not None:
        _insert_before_sect(doc, toc_element)
    else:
        # Fallback: einfaches TOC-Feld
        p = doc.add_paragraph()
        r = p.add_run('Inhaltsverzeichnis – in Word mit F9 aktualisieren')
        r.italic = True
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    _add_horizontal_rule(doc)

    # ── 3. Inhalt ──────────────────────────────────────────────────────────────
    if markdown_content:
        _add_markdown_content(doc, markdown_content, style_map, fallback_color,
                              skip_first_h1=True)
    else:
        sec_num = 1

        if aushang and template_id == 'intern':
            _add_heading_safe(doc, f'{sec_num}. Aushang', 2, style_map, fallback_color)
            doc.add_paragraph('Dieses Dokument hängt/liegt im Bereich aus.')
            _add_warning_box(doc,
                '⚠️  Die ausgedruckte Version kann evtl. einen veralteten Stand haben '
                '– gültig ist immer die digitale Version im WBI-System.')
            _add_horizontal_rule(doc)
            sec_num += 1

        if refs:
            _add_heading_safe(doc, f'{sec_num}. Mitgeltende Unterlagen', 2, style_map, fallback_color)
            _add_refs_table(doc, refs, style_map)
            _add_horizontal_rule(doc)
            sec_num += 1

        valid_chapters = [c for c in chapters if c.get('name', '').strip()]
        if valid_chapters:
            for idx, ch in enumerate(valid_chapters):
                n = sec_num + idx
                _add_heading_safe(doc, f'{n}. {ch["name"]}', 2, style_map, fallback_color)
                subs = [s for s in ch.get('subs', []) if s.strip()]
                if subs:
                    for si, sub in enumerate(subs):
                        _add_heading_safe(doc, f'{n}.{si+1} {sub}', 3, style_map, fallback_color)
                        doc.add_paragraph('Inhalt hier einfügen.')
                else:
                    doc.add_paragraph('Inhalt hier einfügen.')
                if idx < len(valid_chapters) - 1:
                    _add_horizontal_rule(doc)
        else:
            _add_heading_safe(doc, f'{sec_num}. Kapitel', 2, style_map, fallback_color)
            doc.add_paragraph('Inhalt hier einfügen.')

    buf = io.BytesIO()
    doc.save(buf)
    return buf, _safe_filename(title)
